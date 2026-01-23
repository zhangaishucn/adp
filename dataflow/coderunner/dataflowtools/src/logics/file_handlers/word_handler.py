#!/usr/bin/python3
# -*- coding:utf-8 -*-

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.opc.exceptions import PackageNotFoundError
from .base_handler import BaseFileHandler
from models.file_operator import FileOperatorRequest
from common.configs import InsertType


class DocxHandler(BaseFileHandler):
    """DOCX文件处理器"""
    
    SUPPORT_INSERT_TYPES = [
        InsertType.APPEND.value,
        InsertType.APPEND_BEFORE.value,
        InsertType.APPEND_AFTER.value,
        InsertType.COVER.value,
    ]

    def get_file_extension(self) -> str:
        return '.docx'
    
    def needs_temp_file(self) -> bool:
        return True
    
    def supports_operation(self, request: FileOperatorRequest) -> bool:
        """检查是否支持该操作"""
        # DOCX只支持文本内容插入，不支持行列操作
        return request.new_type is None and request.insert_type in self.SUPPORT_INSERT_TYPES
    
    def _do_update(self, source_path: str, target_path: str, request: FileOperatorRequest) -> None:
        """
        统一的更新接口
        
        Args:
            file_path: 临时文件路径
            request: 更新请求
            
        Returns:
            更新后的数据
        """
        
        try:
            doc = Document(source_path)
        except PackageNotFoundError:
            doc = Document()
        
        # 处理文档内容
        if request.insert_type == InsertType.COVER.value:
            self._clear_document(doc)
        
        # 插入内容
        self._insert_content(doc, request)
        
        doc.save(target_path)
    
    def _insert_content(self, doc: Document, request: FileOperatorRequest):
        """插入内容"""
        content = request.content
        
        # 如果content是字符串，直接插入
        if isinstance(content, str):
            self._add_formatted_paragraph(doc, content)
        # 如果是列表，逐个插入
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    self._add_formatted_paragraph(doc, item)
                else:
                    self._add_formatted_paragraph(doc, str(item))
        else:
            self._add_formatted_paragraph(doc, str(content))
    
    @staticmethod
    def _clear_document(doc: Document):
        """清除文档所有内容"""
        for element in list(doc.element.body):
            doc.element.body.remove(element)
    
    @staticmethod
    def _add_formatted_paragraph(doc: Document, text: str):
        """添加格式化段落"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        # 设置字体样式
        font = run.font
        font.name = '宋体'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # 设置东亚语言字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')